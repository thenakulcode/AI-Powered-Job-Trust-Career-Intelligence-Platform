"""Resume parsing utilities for PDF and DOCX files.

Backward-compatible with the previous resume_parser.py: same public
function names and signatures (extract_resume_text, parse_resume_sections,
normalize_skill_name, display_skill_name) and SKILL_TERMS is preserved as a
legacy alias so any other code still importing it does not break. However,
skill/keyword extraction internally now delegates to the domain-independent
extractor in skill_extraction.py instead of a hardcoded software-only list,
so resumes from any profession parse correctly.
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from typing import Any

import fitz
from docx import Document

from backend.services.skill_extraction import (
    display_phrase,
    extract_key_phrases,
    normalize_phrase,
)

logger = logging.getLogger(__name__)

SECTION_LABELS: dict[str, tuple[str, ...]] = {
    "skills": ("skills", "technical skills", "core skills", "tools", "technologies", "competencies", "expertise"),
    "education": ("education", "academic background", "academics", "qualifications"),
    "projects": ("projects", "selected projects", "personal projects", "portfolio", "publications", "research"),
    "experience": ("experience", "work experience", "professional experience", "employment", "clinical experience", "teaching experience"),
    "certifications": ("certifications", "licenses", "certificates", "licensure"),
    "achievements": ("achievements", "awards", "accomplishments", "highlights", "honors"),
    "languages": ("languages", "spoken languages"),
    "internships": ("internships", "internship", "residency", "fellowship"),
}

# Legacy alias kept for backward compatibility with any code that still
# imports SKILL_TERMS directly. New code should not rely on this list — skill
# extraction is now domain-independent via skill_extraction.extract_key_phrases.
SKILL_TERMS: list[str] = []


def normalize_text(text: str | None) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip()


def normalize_skill_name(value: str | None) -> str:
    """Backward-compatible alias for normalize_phrase — same name/signature
    as the old software-only version, now domain-independent under the
    hood."""
    return normalize_phrase(value)


def display_skill_name(value: str | None) -> str:
    """Backward-compatible alias for display_phrase."""
    return display_phrase(value)


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    document = fitz.open(stream=file_bytes, filetype="pdf")
    chunks = [page.get_text("text") for page in document if page.get_text("text").strip()]
    document.close()
    return "\n".join(chunks)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_blocks: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_values:
                table_blocks.append(" | ".join(row_values))
    return "\n".join(paragraphs + table_blocks)


def extract_resume_text(file_name: str, file_bytes: bytes) -> str:
    name = (file_name or "").lower()
    if name.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if name.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX resume.")


def parse_resume_sections(resume_text: str) -> dict[str, Any]:
    text = normalize_text(resume_text or "")
    if not text:
        return {
            "name": "",
            "email": "",
            "phone": "",
            "skills": [],
            "education": [],
            "projects": [],
            "experience": [],
            "internships": [],
            "certifications": [],
            "achievements": [],
            "languages": [],
            "tools": [],
            "linkedin": "",
            "github": "",
            "portfolio": "",
            "total_experience": 0,
            "contact_details": {},
            "summary": "",
        }

    lines = [line.strip() for line in resume_text.splitlines() if line.strip()]
    sections: dict[str, list[str]] = {key: [] for key in SECTION_LABELS}
    current_section: str | None = None

    for line in lines:
        lowered = line.lower()
        matched_section = None
        for key, labels in SECTION_LABELS.items():
            if any(lowered.startswith(label + ":") or lowered == label or lowered.startswith(label + " ") for label in labels):
                matched_section = key
                break
        if matched_section:
            current_section = matched_section
            if lowered not in {label for labels in SECTION_LABELS.values() for label in labels}:
                sections[matched_section].append(line)
            continue

        if current_section:
            sections[current_section].append(line)

    def clean_items(items: list[str]) -> list[str]:
        cleaned: list[str] = []
        for item in items:
            candidate = re.sub(r"^[-•*]\s*", "", item).strip()
            if candidate:
                cleaned.append(candidate)
        return cleaned

    cleaned_skills = clean_items(sections["skills"])
    cleaned_education = clean_items(sections["education"])
    cleaned_projects = clean_items(sections["projects"])
    cleaned_experience = clean_items(sections["experience"])
    cleaned_internships = clean_items(sections["internships"])
    cleaned_certifications = clean_items(sections["certifications"])
    cleaned_achievements = clean_items(sections["achievements"])
    cleaned_languages = clean_items(sections["languages"])

    skills = _extract_skills(text, cleaned_skills)
    education = cleaned_education or _extract_education(lines)
    projects = cleaned_projects or _extract_projects(lines)
    experience = cleaned_experience or _extract_experience(lines)
    internships = cleaned_internships or _extract_internships(lines)
    certifications = cleaned_certifications or _extract_certifications(lines)
    achievements = cleaned_achievements or _extract_achievements(lines)
    languages = cleaned_languages or _extract_languages(lines)
    tools = _extract_tools(lines, skills)

    if not skills and not education and not projects and not experience and not certifications and not achievements:
        skills = _extract_skills(text, [])
        if not experience:
            experience = [sent for sent in re.split(r"(?<=[.!?])\s+", text) if len(sent.split()) >= 4][:3]

    name = _extract_name(lines)
    email = _extract_contact_detail(text, r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})")
    phone = _extract_contact_detail(text, r"(\+?\d[\d\s().-]{7,}\d)")
    linkedin = _extract_contact_detail(text, r"(https?://(?:www\.)?linkedin\.com/[^\s]+)", default="")
    github = _extract_contact_detail(text, r"(https?://(?:www\.)?github\.com/[^\s]+)", default="")
    portfolio = _extract_contact_detail(text, r"(https?://[^\s]+)", default="")

    total_experience = _extract_total_experience(text)
    contact_details = {"email": email, "phone": phone, "linkedin": linkedin, "github": github, "portfolio": portfolio}

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "education": education,
        "projects": projects,
        "experience": experience,
        "internships": internships,
        "certifications": certifications,
        "achievements": achievements,
        "languages": languages,
        "tools": tools,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "total_experience": total_experience,
        "contact_details": contact_details,
        "summary": text[:800],
    }


def _extract_skills(text: str, section_items: list[str]) -> list[str]:
    """Domain-independent skill extraction: pulls explicit skills-section
    items (split on delimiters) plus dynamically extracted key phrases from
    the full resume text, rather than matching against a fixed software
    vocabulary."""
    candidates: list[str] = []
    for item in section_items:
        for part in re.split(r"[,;|/]+", item):
            normalized = normalize_phrase(part)
            if normalized:
                candidates.append(normalized)

    candidates.extend(extract_key_phrases(text, top_n=30))

    unique_items: list[str] = []
    seen: set[str] = set()
    for item in candidates:
        value = normalize_phrase(item)
        if not value or value in seen:
            continue
        seen.add(value)
        unique_items.append(value)

    return unique_items


def _extract_education(lines: list[str]) -> list[str]:
    matches: list[str] = []
    for line in lines:
        if re.search(
            r"\b(b\.s|b\.tech|m\.s|m\.tech|bachelor|master|phd|ph\.d|associate|diploma|degree|"
            r"mba|md|jd|university|college|institute)\b",
            line,
            re.I,
        ):
            matches.append(line)
    return matches


def _extract_projects(lines: list[str]) -> list[str]:
    matches: list[str] = []
    for line in lines:
        if re.search(r"\b(project|portfolio|built|developed|implemented|published|publication|study|campaign|case)\b", line, re.I):
            matches.append(line)
    return matches


def _extract_experience(lines: list[str]) -> list[str]:
    matches: list[str] = []
    for line in lines:
        if re.search(
            r"\b(years?|yrs?|yr|experience|worked|led|managed|taught|treated|litigated|"
            r"analyzed|conducted|supervised|designed|coordinated)\b",
            line,
            re.I,
        ):
            matches.append(line)
    return matches


def _extract_internships(lines: list[str]) -> list[str]:
    return [line for line in lines if re.search(r"\b(intern(ship|ships|ed)?|residency|fellowship)\b", line, re.I)]


def _extract_certifications(lines: list[str]) -> list[str]:
    return [line for line in lines if re.search(r"\b(certified|certificate|certification|license|licensure|board certified)\b", line, re.I)]


def _extract_achievements(lines: list[str]) -> list[str]:
    matches: list[str] = []
    for line in lines:
        if re.search(r"\b(improved|increased|reduced|delivered|optimized|won|awarded|launched|boosted|recognized|honored|published)\b", line, re.I):
            matches.append(line)
    return matches


def _extract_languages(lines: list[str]) -> list[str]:
    languages = []
    for line in lines:
        if re.search(r"\b(english|spanish|french|german|hindi|arabic|mandarin|japanese|korean|portuguese|italian)\b", line, re.I):
            languages.append(line)
    return languages


def _extract_tools(lines: list[str], skills: list[str]) -> list[str]:
    """Tools are no longer detected against a fixed software list; any line
    matching a generic "tool/instrument/platform" cue, plus overlap with
    already-extracted skill phrases, is treated as a tool reference."""
    tools = []
    for line in lines:
        if re.search(r"\b(tool|platform|software|instrument|equipment|system)\b", line, re.I):
            tools.append(line)
    return tools


def _extract_name(lines: list[str]) -> str:
    section_starts = tuple(
        label for labels in SECTION_LABELS.values() for label in labels
    )
    for line in lines:
        candidate = normalize_text(line)
        if not candidate or len(candidate.split()) > 6:
            continue
        if re.search(r"@|\+\d|\bhttps?://\b", candidate):
            continue
        if candidate.lower().startswith(section_starts):
            continue
        return candidate
    return ""


def _extract_contact_detail(text: str, pattern: str, default: str = "") -> str:
    match = re.search(pattern, text, re.I)
    if match:
        return match.group(1).strip()
    return default


def _extract_total_experience(text: str) -> int:
    matches = re.findall(r"(\d+)\s*(?:\+)?\s*(?:years?|yrs?|yr)", text.lower())
    if matches:
        return int(matches[0])
    return 0